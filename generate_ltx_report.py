#!/usr/bin/env python3
"""Generate LTX Trading Site Analysis Report as Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Heading styles
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0, 32, 96)

def add_table_with_style(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_screenshot(doc, path, width=5.5, caption=""):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        if caption:
            p = doc.add_paragraph(caption)
            p.style = doc.styles['Normal']
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(8)

# ====== TITLE PAGE ======
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('LTX Trading Website\nSite Analysis Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
subtitle = doc.add_paragraph('Comprehensive Migration Assessment')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.color.rgb = RGBColor(0, 32, 96)
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Website: ').bold = True
info.add_run('https://www.ltxtrading.com/\n')
info.add_run('Date: ').bold = True
info.add_run('March 5, 2026\n')
info.add_run('Total Pages in Sitemap: ').bold = True
info.add_run('91 URLs\n')
info.add_run('Current Platform: ').bold = True
info.add_run('Custom (jQuery-based, formerly WordPress/Divi)\n')
doc.add_page_break()

# ====== TABLE OF CONTENTS ======
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Templates Inventory',
    '3. Blocks / Components Catalog',
    '4. Page Counts by Template',
    '5. Integrations Analysis',
    '6. Complex Use Cases & Observations',
    '7. Migration Estimates',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ====== 1. EXECUTIVE SUMMARY ======
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'LTX Trading (www.ltxtrading.com) is a corporate bond e-trading platform owned by Broadridge '
    'Financial Solutions. The website serves as a marketing and lead generation platform for their '
    'AI-powered bond trading technology products including BondGPT, Liquidity Cloud, RFX Trading Protocol, and RFQ+.'
)
doc.add_paragraph(
    'The site contains 91 URLs in its sitemap, built on a custom platform with jQuery and remnants '
    'of WordPress/Divi theme architecture (evidenced by et_pb_* JavaScript references). The site uses '
    'a modular section-based layout with consistent blocks across pages. The relatively small size and '
    'standardized templates make it well-suited for migration to Adobe Edge Delivery Services.'
)

key_findings = [
    '91 total pages identified in the sitemap',
    '8 distinct page templates identified',
    '15+ reusable blocks/components cataloged',
    'Custom jQuery-based CMS with WordPress/Divi heritage',
    'Lead generation forms on most pages (custom implementation)',
    'Third-party integrations: Google Tag Manager, OneTrust (cookie consent), The Trade Desk, ZoomInfo, Clickagy, Adobe Fonts (Typekit), Google Fonts',
    'Majority of pages follow standardized templates suitable for automated migration',
]
doc.add_heading('Key Findings:', level=3)
for f in key_findings:
    doc.add_paragraph(f, style='List Bullet')
doc.add_page_break()

# ====== 2. TEMPLATES INVENTORY ======
doc.add_heading('2. Templates Inventory', level=1)
doc.add_paragraph(
    'The following unique page templates were identified across the LTX Trading website. '
    'Templates are categorized by their structural layout and content organization patterns.'
)

templates = [
    ['Homepage', 'High',
     'Unique layout with hero, BondGPT promo banner, icon cards, product cards (buy-side/dealers), '
     'CTA band, quote, text cards (4 products), content cards (3 articles), two-column stats/text, '
     'and demo form. Most complex page with the most diverse block composition.',
     'https://www.ltxtrading.com/'],
    ['Audience Page\n(Dealers / Buy Side)', 'Medium',
     'Hero with CTA, image carousel with text slides, vertical tabs/accordion block with 5 features, '
     'CTA band, quote block, 3 content cards, and demo request form. Reused for Dealers and Buy Side pages.',
     'https://www.ltxtrading.com/dealers\nhttps://www.ltxtrading.com/buy-side'],
    ['Product Page\n(Platform Products)', 'Medium-High',
     'Hero with CTA, detailed content sections with lists, CTA bands, carousels (some with video embeds), '
     'two-column feature comparisons, interactive FAQ/micro-survey sections, demo form, and content cards. '
     'Each product page has unique variations.',
     'https://www.ltxtrading.com/bondgpt\nhttps://www.ltxtrading.com/liquidity-cloud\n'
     'https://www.ltxtrading.com/rfx-trading-protocol\nhttps://www.ltxtrading.com/rfqplus'],
    ['Insights Hub\n(Listing Page)', 'Medium',
     'Hero with CTA, categorized card grid (Featured, Videos, More Insights) displaying content cards '
     'with type badges (VIDEO, REPORT, ARTICLE, RESOURCE, WEBINAR, PODCAST, WHITEPAPER). Extensive listing '
     'with 30+ items organized in 3-column rows. Includes demo form.',
     'https://www.ltxtrading.com/insights'],
    ['News Listing Page', 'Medium',
     'Hero without CTA, chronological list of news items with headline, date, optional excerpt, '
     'and "READ MORE" links. Simple repeating card structure, 40+ items. No demo form. '
     'Mix of internal articles and external links.',
     'https://www.ltxtrading.com/news'],
    ['About Us Page', 'High',
     'Hero with CTA, product overview with two audience cards, CTA band, text navigation cards (3), '
     'statistics/metrics section with Broadridge data, awards carousel with images, '
     'Leadership Team section with expandable bios and LinkedIn links, '
     'Board of Directors section with expandable bios. Unique template.',
     'https://www.ltxtrading.com/about-us'],
    ['Article / Press Release\nPage', 'Low',
     'Simple layout with header/nav, article title (H1), optional subtitle (H2 italic), '
     'body text with paragraphs and inline links, and footer. No hero block, no form, no sidebar. '
     'Clean editorial template used for press releases and articles.',
     'https://www.ltxtrading.com/ltx-expands-genai-capabilities-with-bondgpt-intelligence\n'
     'https://www.ltxtrading.com/broadridge-announces-new-patent-on-large-language-model-orchestration-of-machine'],
    ['Contact / Form Page', 'Low',
     'Hero with text, contact form (same as other pages but primary focus), '
     'corporate address/phone details, and footer. Minimal content sections.',
     'https://www.ltxtrading.com/contact-us'],
]

headers = ['Template Name', 'Complexity', 'Description & Reasoning', 'Reference URL(s)']
add_table_with_style(doc, headers, templates)

doc.add_paragraph()
doc.add_heading('Template Screenshots', level=2)

screenshot_map = {
    'Homepage': 'screenshots/homepage-full.png',
    'Dealers (Audience Page)': 'screenshots/dealers.png',
    'Buy Side (Audience Page)': 'screenshots/buy-side.png',
    'BondGPT (Product Page)': 'screenshots/bondgpt.png',
    'Insights Hub (Listing)': 'screenshots/insights.jpeg',
    'About Us': 'screenshots/about-us.jpeg',
    'News (Listing)': 'screenshots/news.png',
    'Article / Press Release': 'screenshots/article-press-release.png',
    'Contact Us (Form)': 'screenshots/contact-us.png',
}

for name, path in screenshot_map.items():
    full_path = os.path.join('/workspace', path)
    if os.path.exists(full_path):
        doc.add_heading(name, level=3)
        try:
            doc.add_picture(full_path, width=Inches(4.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'[Screenshot: {path} - Error: {e}]')

doc.add_page_break()

# ====== 3. BLOCKS / COMPONENTS CATALOG ======
doc.add_heading('3. Blocks / Components Catalog', level=1)
doc.add_paragraph(
    'The following reusable blocks and components were identified across the site. '
    'Components with the same content model but different visual treatments are noted as design variations.'
)

blocks = [
    ['Header / Navigation', 'Medium',
     'Sticky top navigation with LTX logo, hamburger menu for mobile, main nav links '
     '(Home, Dealers, Buy Side, Platform [dropdown], Insights, About [dropdown], Contact Us). '
     'Dropdown menus for Platform and About sections. Consistent across all pages.',
     'All pages'],
    ['Hero Banner', 'Low',
     'Full-width gradient background (dark blue/purple with particle animation), H1 heading, '
     'subtitle paragraph, and optional CTA button ("Request a demo"). '
     'Design variation: some heroes have CTA, some don\'t (News page). All use same layout pattern.',
     'All main pages'],
    ['Carousel / Slider Block', 'Medium',
     'Content carousel with text slides, navigation arrows (prev/next), and dot indicators. '
     'Used on Dealers and Buy Side pages with 3 text-based slides. '
     'BondGPT page has a variation with video embeds in carousel slides.',
     'https://www.ltxtrading.com/dealers\nhttps://www.ltxtrading.com/buy-side\nhttps://www.ltxtrading.com/bondgpt'],
    ['Vertical Tabs / Accordion', 'Medium',
     'Left-aligned tab labels with expandable content areas on the right. 5 tabs showing feature details. '
     'Clean accordion-style interaction. Used on audience pages to present platform features.',
     'https://www.ltxtrading.com/dealers\nhttps://www.ltxtrading.com/buy-side'],
    ['CTA Band', 'Low',
     'Full-width gradient background strip with centered H2 heading and CTA button. '
     'Used as section dividers to drive demo requests. Variants: "Demo LTX" and "Request a BondGPT demo".',
     'https://www.ltxtrading.com/\nhttps://www.ltxtrading.com/dealers\nhttps://www.ltxtrading.com/buy-side\nhttps://www.ltxtrading.com/bondgpt'],
    ['Quote / Testimonial Block', 'Low',
     'Centered blockquote with italic text on gradient background, attribution line below. '
     'Simple design with decorative quotation marks.',
     'https://www.ltxtrading.com/\nhttps://www.ltxtrading.com/dealers\nhttps://www.ltxtrading.com/buy-side'],
    ['Content Cards (3-up)', 'Low',
     '3-column grid of linked cards, each with background image, type badge '
     '(VIDEO, PRESS RELEASE, ARTICLE, RESOURCE, REPORT, etc.), title text, and ">>" arrow. '
     'Cards link to articles, PDFs, or external URLs. Used across most main pages.',
     'https://www.ltxtrading.com/\nhttps://www.ltxtrading.com/dealers\nhttps://www.ltxtrading.com/buy-side\nhttps://www.ltxtrading.com/bondgpt\nhttps://www.ltxtrading.com/insights'],
    ['Text Cards (Product Nav)', 'Low',
     '4-column grid of cards linking to product pages. Each card has product name (with ®/SM marks), '
     'description text, and ">>" arrow link. Used on homepage to navigate to platform products.',
     'https://www.ltxtrading.com/'],
    ['Icon Cards', 'Low',
     'Section with H2 heading, paragraph, and 3 icon+text items in a row '
     '(Discover liquidity, Know when to trade, Execute like never before). Uses SVG/icon imagery.',
     'https://www.ltxtrading.com/'],
    ['Product Module (2-card)', 'Low',
     'Two side-by-side cards (Buy Side / Dealers) with H3, description, and link button. '
     'Used to direct visitors to audience-specific pages.',
     'https://www.ltxtrading.com/\nhttps://www.ltxtrading.com/about-us'],
    ['Two-Column Text + Stats', 'Medium',
     'Left column: paragraph text with "Learn more" link. Right column: bullet list with statistics '
     '(products supported, average sizes, advisory group members, ownership). Used on homepage.',
     'https://www.ltxtrading.com/'],
    ['Demo Request Form', 'Medium',
     'Contact/lead generation form with fields: First Name, Last Name, Company Name, Email, Phone, '
     'Country (dropdown with 200+ options), and textarea for inquiry description. Submit button. '
     'Custom JavaScript validation (ltx-form.js). Appears on most main pages.',
     'https://www.ltxtrading.com/\nhttps://www.ltxtrading.com/dealers\nhttps://www.ltxtrading.com/buy-side\nhttps://www.ltxtrading.com/bondgpt\nhttps://www.ltxtrading.com/contact-us'],
    ['BondGPT Promo Banner', 'Low',
     'Full-width dark section with H2 text about BondGPT and "Discover BondGPT" CTA button. '
     'Variation of CTA Band specific to BondGPT promotion.',
     'https://www.ltxtrading.com/'],
    ['Micro Survey / FAQ Block', 'High',
     'Interactive Q&A component with question list on the left and answer panel on the right. '
     'Clicking a question reveals the answer with "Go back" navigation. '
     'Includes links to PDF resources. Complex JavaScript interaction.',
     'https://www.ltxtrading.com/bondgpt'],
    ['Leadership / Team Block', 'Medium',
     'Grid of team member cards with name (H4), title, truncated bio with "Show More" toggle, '
     'and LinkedIn profile link icon. Used for Leadership Team and Board of Directors sections.',
     'https://www.ltxtrading.com/about-us'],
    ['Awards Carousel', 'Medium',
     'Grid display of award badges/logos with award name and year. 2 rows of 3 awards each. '
     'Each award links to external source article.',
     'https://www.ltxtrading.com/about-us'],
    ['Statistics / Metrics Block', 'Low',
     'Left: H2 heading + paragraph about Broadridge. Right: 3 large stat numbers with descriptions '
     '($6B+ revenues, $10T daily securities, 20 of 25 primary dealers).',
     'https://www.ltxtrading.com/about-us'],
    ['News List Block', 'Medium',
     'Chronological list of news items. Each item: H4 headline, date paragraph, optional excerpt, '
     '"READ MORE" link. Alternating internal/external links. 40+ items without pagination.',
     'https://www.ltxtrading.com/news'],
    ['Footer', 'Medium',
     'Three-column footer with: (1) Social icons (LinkedIn, Twitter/X), (2-3) Navigation links organized '
     'in two columns matching main nav. Below: copyright, legal links (Terms of Use, Accessibility, '
     'Legal Statements, Privacy, Do Not Sell, Your Privacy Choices). Full-width disclaimer paragraph. '
     'Consistent across all pages.',
     'All pages'],
]

headers = ['Block Name', 'Complexity', 'Description & Behaviour', 'Reference URL(s)']
add_table_with_style(doc, headers, blocks)

doc.add_page_break()

# ====== 4. PAGE COUNTS BY TEMPLATE ======
doc.add_heading('4. Page Counts by Template', level=1)
doc.add_paragraph(
    'Based on analysis of the 91 URLs in the sitemap and the identified template patterns:'
)

page_counts = [
    ['Homepage', '1', 'Automatic',
     'Standardized modular layout. Requires custom block styling but content is static.'],
    ['Audience Page\n(Dealers / Buy Side)', '2', 'Automatic',
     'Identical template structure, different content. Carousel and vertical tabs reusable.'],
    ['Product Page\n(BondGPT, Liquidity Cloud,\nRFX, RFQ+)', '4-6', 'Semi-Automatic',
     'Similar structure but each product page has unique sections and interactive elements. '
     'BondGPT page has video carousel and micro-survey requiring custom development.'],
    ['Insights Hub', '1', 'Manual',
     'Large listing page with 30+ categorized content cards. Cards link to mix of '
     'internal pages, external URLs, and PDF documents. Content card categorization needs manual mapping.'],
    ['News Listing', '1', 'Manual',
     '40+ news items mixing internal press releases and external article links. '
     'No pagination. Chronological ordering with dates.'],
    ['About Us', '1', 'Manual',
     'Unique layout with leadership bios (expandable), board of directors, awards, '
     'statistics. Interactive "Show More" toggles need custom JS.'],
    ['Article / Press Release', '~30', 'Automatic',
     'Simple article template with title, body text, and footer. '
     'Standardized layout, content-only differences. Easily automated.'],
    ['Content Pages\n(Misc: careers, partners,\nsolution, ask-bondgpt,\nwhitepapers, videos, etc.)', '~15-20', 'Semi-Automatic',
     'Various content pages with mixed layouts. Some are simple text pages, '
     'others have embedded videos (YouTube/Vimeo), PDF links, or unique interactive elements. '
     'The "Ask BondGPT" and "Careers" pages may have unique functionality.'],
    ['Utility Pages\n(404, sitemap, preview pages)', '3-5', 'Automatic',
     'Standard utility pages. 404 error page, HTML sitemap, and preview/staging pages.'],
    ['TOTAL', '~58-67\n(excl. duplicates\n& utility)', '—', '91 URLs in sitemap including some duplicates and utility pages.'],
]

headers = ['Template Type', 'Page Count', 'Migration Type', 'Notes']
add_table_with_style(doc, headers, page_counts)

doc.add_paragraph()
doc.add_heading('Migration Approach Summary', level=2)
migration_summary = [
    ['Automatic Migration', '~35-38', 'Article/press release pages, audience pages, homepage, simple content pages. '
     'Standardized templates with predictable block composition.'],
    ['Semi-Automatic Migration', '~15-20', 'Product pages and miscellaneous content pages requiring some manual '
     'block mapping and content adjustment. Video embeds and interactive elements need review.'],
    ['Manual Migration', '~8-10', 'Insights hub, news listing, about us, careers, and pages with unique '
     'interactive functionality (micro-surveys, expandable bios, etc.).'],
]
headers = ['Approach', 'Est. Pages', 'Description']
add_table_with_style(doc, headers, migration_summary)
doc.add_page_break()

# ====== 5. INTEGRATIONS ANALYSIS ======
doc.add_heading('5. Integrations Analysis', level=1)
doc.add_paragraph(
    'The following third-party integrations and embedded services were identified through '
    'analysis of script tags, network requests, meta tags, and page source code.'
)

integrations = [
    ['Google Tag Manager (GTM)', 'Tag Manager', 'Script tag', 'Medium',
     'GTM container ID: GTM-MPSB6HN. Manages analytics and marketing tags. '
     'window.dataLayer is present.',
     'All pages'],
    ['OneTrust (CookieLaw)', 'Cookie Consent', 'Script / SDK', 'Medium',
     'Cookie consent banner and preference center. SDK version 202505.2.0. '
     'Manages GDPR/CCPA compliance with accept/reject/customize options. '
     'Links to Broadridge privacy policies.',
     'All pages'],
    ['The Trade Desk', 'Advertising / Tracking', 'Script / Pixel', 'Low',
     'Advertiser ID: 7ts8ctf. Ad tracking pixel via adsrvr.org. '
     'Includes cookie sync iframe (insight.adsrvr.org). Used for programmatic advertising.',
     'All pages'],
    ['ZoomInfo (zi-scripts)', 'Lead Intelligence', 'Script tag', 'Medium',
     'ZoomInfo tag (zi-tag.js) for B2B visitor identification and lead enrichment. '
     'Tracks website visitors for sales intelligence.',
     'All pages'],
    ['Clickagy', 'Audience Intelligence', 'Script tag', 'Low',
     'Behavioral audience data collection via tags.clickagy.com. '
     'Provides intent data for B2B marketing.',
     'All pages'],
    ['Google Fonts', 'Font Service', 'CSS link', 'Low',
     'Open Sans font family (300-800 weights, italic variants, latin + latin-ext subsets). '
     'Also Cormorant Upright font.',
     'All pages'],
    ['Adobe Fonts (Typekit)', 'Font Service', 'CSS link', 'Low',
     'Typekit ID: aax8mol. Custom web fonts served via use.typekit.net.',
     'All pages'],
    ['Custom Lead Form', 'Form Processing', 'Custom Code', 'High',
     'Custom-built lead capture form (ltx-form.js, 2200+ lines). '
     'Handles validation, country dropdown, and submission to backend API. '
     'Not a standard form service (HubSpot, Marketo, etc.) — appears to be custom-built. '
     'Form endpoint and processing logic need investigation during migration.',
     'Homepage, Dealers, Buy Side, BondGPT, Contact Us, Insights'],
    ['jQuery + jQuery Migrate', 'JS Framework', 'Script tag', 'Low',
     'jQuery core library with jQuery Migrate 1.4 for backward compatibility. '
     'Foundation of all interactive features.',
     'All pages'],
    ['WordPress MediaElement', 'Media Player', 'Script / CSS', 'Low',
     'MediaElement.js player (v4.2.6) with WordPress integration. '
     'Handles video/audio playback on pages with embedded media.',
     'Pages with video content'],
    ['YouTube / Video Embeds', 'Video Hosting', 'Iframe embed', 'Low',
     'Videos embedded via iframes on certain pages (BondGPT carousel, TraderTV pages). '
     'Standard YouTube/video embed integration.',
     'https://www.ltxtrading.com/bondgpt\nTraderTV pages'],
    ['On24 (Webinar Platform)', 'Webinar / Events', 'External link', 'Low',
     'Links to On24 webinar platform (event.on24.com) for recorded webinars. '
     'External redirect, not embedded.',
     'https://www.ltxtrading.com/insights'],
    ['LinkedIn / Twitter (X)', 'Social Media', 'External links', 'Low',
     'Social media profile links in footer and leadership bios. No embedded widgets or share buttons. '
     'LinkedIn: /company/ltxtrading/ | Twitter: @LTXtrading',
     'All pages (footer), About Us (leader bios)'],
    ['Broadridge Legal Hub', 'Legal / Compliance', 'External links', 'Low',
     'Links to broadridge.com for Terms of Use, Privacy Statement, Accessibility, '
     'Legal Statements, Do Not Sell, and Your Privacy Choices pages.',
     'All pages (footer)'],
]

headers = ['Integration', 'Type', 'Method', 'Complexity', 'Description', 'Pages Used']
add_table_with_style(doc, headers, integrations)
doc.add_page_break()

# ====== 6. COMPLEX USE CASES & OBSERVATIONS ======
doc.add_heading('6. Complex Use Cases & Observations', level=1)
doc.add_paragraph(
    'The following complex behaviours, edge cases, and functionality requiring special attention '
    'during migration were identified:'
)

complex_cases = [
    ['Custom Lead Generation Form', '~8 pages',
     'Homepage, Dealers, Buy Side, BondGPT, Liquidity Cloud, RFX, RFQ+, Contact Us, Insights',
     'Custom form implementation (ltx-form.js, 2200+ lines) with client-side validation, '
     'country dropdown with 200+ countries, custom error handling, and backend API integration. '
     'Not built on a standard form service. Form processing endpoint and data flow need reverse-engineering '
     'during migration. This is the primary conversion point for the website.'],
    ['Interactive Micro-Survey\n(Ask BondGPT)', '1 page',
     'https://www.ltxtrading.com/bondgpt',
     'Multi-step FAQ/Q&A interaction where users select questions and see answers with "Go back" navigation. '
     '5 question-answer pairs with links to PDF resources. Requires custom JavaScript development '
     'to replicate the show/hide panel interaction.'],
    ['Expandable Leadership Bios', '1 page',
     'https://www.ltxtrading.com/about-us',
     '13 leadership/board member profiles with "Show More" toggle to expand truncated bios. '
     'Each profile includes name, title, bio text, and LinkedIn link. '
     'Requires custom JS toggle functionality in EDS.'],
    ['Carousel / Slider Components', '3-4 pages',
     'Dealers, Buy Side, BondGPT',
     'Custom carousel implementations with text slides, navigation arrows, and dot indicators. '
     'BondGPT page has a variation with embedded video in carousel slides. '
     'EDS does not have a native carousel block — requires custom block development.'],
    ['Vertical Tabs Component', '2 pages',
     'Dealers, Buy Side',
     'Interactive tabbed content with 5 tabs showing feature descriptions. '
     'Click-to-reveal interaction pattern. Requires custom block development in EDS.'],
    ['Video Embed Pages', '~5-7 pages',
     'TraderTV pages, BondGPT, various article pages',
     'Pages with embedded video players (YouTube iframes, MediaElement.js). '
     'Some are in carousels. Need to ensure proper responsive video embed handling.'],
    ['Mixed Internal/External\nLink Handling', '2 pages',
     'News, Insights',
     'News and Insights listing pages contain cards linking to both internal pages AND external URLs '
     '(Forbes, Bloomberg, CNBC, Markets Media, etc.) as well as PDF documents. '
     'Link routing and "target=_blank" behavior needs careful handling.'],
    ['PDF Document Hosting', '~15+ resources',
     'Insights, product pages, various cards',
     'Multiple PDF documents hosted at /assets/docs/ (sell sheets, whitepapers, research reports). '
     'Used as link targets from content cards and inline links. '
     'These assets need to be migrated and URL paths preserved.'],
    ['Legacy WordPress/Divi\nJS References', 'All pages',
     'All pages',
     'Console errors reference et_pb_init_nav_menu and et_pb_box_shadow_apply_overlay — '
     'Divi theme functions that no longer exist. The site was likely migrated from WordPress/Divi '
     'to a custom solution but retains legacy JS dependencies causing console errors. '
     'This is a cleanup opportunity during migration.'],
    ['Particle/Animation Hero\nBackgrounds', '~10+ pages',
     'All main pages with hero blocks',
     'Hero sections use animated particle/wave effects over gradient backgrounds. '
     'Custom CSS animations. May need WebGL or canvas-based implementation '
     'or simplified CSS-only approach for EDS.'],
]

headers = ['Use Case', 'Instances', 'Where Found', 'Why It\'s Complex']
add_table_with_style(doc, headers, complex_cases)
doc.add_page_break()

# ====== 7. MIGRATION ESTIMATES ======
doc.add_heading('7. Migration Estimates', level=1)
doc.add_paragraph(
    'The following estimates are based on the site analysis, assuming migration to Adobe Edge Delivery Services. '
    'Estimates consider the modular architecture, number of unique templates, custom block development needs, '
    'and integration complexity.'
)

doc.add_heading('7.1 Block Development Effort', level=2)
block_effort = [
    ['Header / Navigation', 'Medium', '2-3 days', 'Dropdown menus, mobile hamburger, sticky behavior'],
    ['Hero Banner', 'Low', '1-2 days', 'Gradient backgrounds, particle animation (simplified)'],
    ['Carousel / Slider', 'High', '3-4 days', 'Custom block with text/video slides, nav arrows, dots'],
    ['Vertical Tabs', 'Medium', '2-3 days', 'Tab interaction, mobile accordion fallback'],
    ['CTA Band', 'Low', '0.5 days', 'Simple gradient section with heading and button'],
    ['Quote / Testimonial', 'Low', '0.5 days', 'Styled blockquote'],
    ['Content Cards (3-up)', 'Low', '1-2 days', 'Card grid with type badges and image backgrounds'],
    ['Text Cards (4-up)', 'Low', '1 day', 'Product navigation cards'],
    ['Icon Cards', 'Low', '0.5-1 day', 'Icon + text layout'],
    ['Product Module (2-card)', 'Low', '0.5-1 day', 'Two audience cards'],
    ['Two-Column Text + Stats', 'Low', '1 day', 'Text and bullet list layout'],
    ['Demo Request Form', 'High', '4-5 days', 'Custom form with validation, country dropdown, API integration'],
    ['Micro Survey / FAQ', 'High', '3-4 days', 'Interactive Q&A with show/hide panels'],
    ['Leadership / Team Block', 'Medium', '2-3 days', 'Expandable bios, LinkedIn links'],
    ['Awards Grid', 'Low', '1 day', 'Image grid with links'],
    ['Statistics Block', 'Low', '0.5-1 day', 'Large numbers with descriptions'],
    ['News List Block', 'Medium', '2-3 days', 'Repeating news items with dates and links'],
    ['Footer', 'Medium', '1-2 days', 'Multi-column nav, legal links, disclaimer'],
    ['TOTAL', '', '25-38 days', 'Block development and styling'],
]
headers = ['Block', 'Complexity', 'Estimated Effort', 'Notes']
add_table_with_style(doc, headers, block_effort)

doc.add_paragraph()
doc.add_heading('7.2 Content Migration Effort', level=2)
content_effort = [
    ['Homepage', '1 page', '2-3 days', 'Complex layout, all blocks combined'],
    ['Audience Pages', '2 pages', '1-2 days', 'Template reuse, content swap'],
    ['Product Pages', '4-6 pages', '4-6 days', 'Each page has unique sections'],
    ['Insights Hub', '1 page', '2-3 days', '30+ content cards to map'],
    ['News Listing', '1 page', '2-3 days', '40+ news items, link verification'],
    ['About Us', '1 page', '2-3 days', 'Leadership bios, awards, stats'],
    ['Article / Press Release Pages', '~30 pages', '3-5 days', 'Bulk migration, simple template'],
    ['Content / Misc Pages', '~15-20 pages', '5-8 days', 'Mixed complexity'],
    ['Contact Us', '1 page', '0.5-1 day', 'Form-focused page'],
    ['TOTAL', '~58-67 pages', '20-34 days', 'Content migration and QA'],
]
headers = ['Content Type', 'Pages', 'Effort', 'Notes']
add_table_with_style(doc, headers, content_effort)

doc.add_paragraph()
doc.add_heading('7.3 Integration Migration Effort', level=2)
integration_effort = [
    ['Google Tag Manager', '1-2 days', 'Re-implement GTM container in EDS header'],
    ['OneTrust Cookie Consent', '1-2 days', 'Integrate OneTrust SDK with EDS'],
    ['Lead Generation Form + API', '5-7 days', 'Reverse-engineer form submission, rebuild in EDS, connect to backend'],
    ['The Trade Desk / Clickagy / ZoomInfo', '1-2 days', 'Re-implement tracking pixels via GTM'],
    ['Font Migration (Google + Adobe)', '0.5-1 day', 'Configure font loading in EDS styles'],
    ['Video Embeds', '1-2 days', 'Ensure responsive iframe handling'],
    ['PDF Asset Migration', '1-2 days', 'Move PDFs to DAM, update links'],
    ['TOTAL', '11-18 days', 'Integration setup and testing'],
]
headers = ['Integration', 'Effort', 'Notes']
add_table_with_style(doc, headers, integration_effort)

doc.add_paragraph()
doc.add_heading('7.4 Overall Migration Summary', level=2)
overall = [
    ['Design System / Global Styles', '5-7 days',
     'Extract design tokens (colors, typography, spacing), set up CSS custom properties, '
     'global styles.css with brand consistency'],
    ['Block Development', '25-38 days',
     '18 unique blocks/components to develop including carousel, tabs, form, FAQ, and leadership blocks'],
    ['Content Migration', '20-34 days',
     '58-67 pages across 8 templates. Mix of automated and manual migration.'],
    ['Integration Setup', '11-18 days',
     'GTM, OneTrust, form backend, tracking pixels, fonts, video embeds, PDF assets'],
    ['QA & Testing', '10-15 days',
     'Cross-browser testing, responsive design verification, form submission testing, '
     'link verification, SEO audit, accessibility review, performance optimization'],
    ['UAT & Revisions', '5-8 days',
     'Client review cycles, feedback implementation, final adjustments'],
    ['TOTAL ESTIMATED EFFORT', '76-120 days',
     'Approximately 15-24 weeks with a team of 2-3 developers. '
     'Can be compressed with parallel workstreams.'],
]
headers = ['Phase', 'Effort (Days)', 'Description']
add_table_with_style(doc, headers, overall)

doc.add_paragraph()
doc.add_heading('7.5 Recommended Team & Timeline', level=2)
doc.add_paragraph(
    'Recommended team composition for efficient execution:'
)
team = [
    'Lead Developer / Architect: 1 FTE — Block development, EDS configuration, integrations',
    'Front-end Developer: 1-2 FTE — Block CSS/JS, content migration, responsive design',
    'Content Migration Specialist: 1 FTE — Content mapping, markdown authoring, QA',
    'QA Engineer: 0.5 FTE — Testing, accessibility, performance validation',
]
for t in team:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'With a team of 2-3 full-time developers working in parallel, the migration can be completed in '
    'approximately 10-16 weeks. The largest risk factor is the custom form implementation, '
    'which requires understanding the existing backend API integration.'
)

doc.add_paragraph()
doc.add_heading('7.6 Risk Factors', level=2)
risks = [
    ['Custom Form Backend', 'High', 'Form processing endpoint and API are custom-built. '
     'Backend integration details may require coordination with the LTX engineering team.'],
    ['Legacy JS Dependencies', 'Low', 'WordPress/Divi remnants cause console errors but don\'t affect functionality. '
     'Migration provides a clean-slate opportunity.'],
    ['PDF Asset Volume', 'Low', '15+ PDF documents need migration. Straightforward but requires URL redirect mapping.'],
    ['External Link Maintenance', 'Medium', 'News and Insights pages link to 30+ external sites. '
     'Links may break over time. Consider link validation during migration.'],
    ['Hero Animations', 'Medium', 'Particle/wave animations in hero backgrounds may need simplification '
     'for EDS. Performance impact should be evaluated.'],
    ['SEO Redirect Mapping', 'Medium', 'All 91 URLs need 301 redirects if URL structure changes. '
     'Maintaining URL structure minimizes this risk.'],
]
headers = ['Risk', 'Impact', 'Description']
add_table_with_style(doc, headers, risks)

# Save
output_path = '/workspace/LTX_Trading_Site_Analysis_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
