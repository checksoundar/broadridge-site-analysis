#!/usr/bin/env python3
"""
Generate Broadridge.com Site Analysis Report as MS Word Document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x5B)  # Broadridge navy

# ── Helper Functions ────────────────────────────────────────────
def add_table_with_data(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A235B"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_val in enumerate(row_data):
            row_cells[c_idx].text = str(cell_val)
            for paragraph in row_cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        # Alternate row shading
        if r_idx % 2 == 0:
            for c_idx in range(len(headers)):
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F2F8"/>')
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    return table


def add_screenshot(doc, img_path, caption, width=5.5):
    """Add a screenshot with caption."""
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style.font.italic = True
        cap.style.font.size = Pt(8)
        cap.style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Broadridge.com')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x5B)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive Site Analysis Report')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph('')

detail = doc.add_paragraph()
detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = detail.add_run('Adobe Edge Delivery Services Migration Assessment')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Prepared: March 2026\nTarget URL: https://www.broadridge.com/')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Templates Inventory',
    '3. Blocks / Components Catalog',
    '4. Page Counts by Template',
    '5. Integrations Analysis',
    '6. Complex Use Cases & Observations',
    '7. Migration Estimates',
    '8. Appendix: Screenshots',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

exec_text = """Broadridge Financial Solutions (broadridge.com) is a large-scale enterprise website serving the global financial services industry. The site encompasses approximately 1,084 indexed pages across 12 distinct page templates, supported by ~25 reusable UI blocks/components.

Key findings from this analysis:

• The site runs on a custom CMS ("Broadridge V2") with jQuery 3.7.1 and Alpine.js as the primary JavaScript frameworks, plus Swiper.js for carousels and Lottie for hero animations.

• A subset of legacy pages (primarily /case-study/ URLs) run on an older CMS platform with a different navigation structure, different form system, and legacy jQuery patterns. These require special handling during migration.

• The site supports 4 locales: English (primary), French (fr/), German (de/), and Japanese (jp/), with a geolocation-based auto-redirect system.

• 16 third-party integrations were identified, including OneTrust (cookie consent), Google Tag Manager, reCAPTCHA, ZoomInfo, Facebook Pixel, LinkedIn Insight, Microsoft Clarity, Bing UET, The Trade Desk, Clickagy, CrownPeak accessibility, and Pardot/Marketing Cloud (pages.broadridge.com).

• The majority of pages (~70%) follow repeatable templates suitable for automated migration, while ~30% contain unique layouts or legacy structures requiring manual attention."""

doc.add_paragraph(exec_text)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2. TEMPLATES INVENTORY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. Templates Inventory', level=1)

doc.add_paragraph('The following 12 unique page templates were identified across the Broadridge.com site. Each template represents a distinct layout pattern with specific block combinations and content structures.')

templates_data = [
    ['Homepage', 'High',
     'Animated Lottie hero, featured solutions cards, insights carousel, industry tabs, capabilities accordion, awards grid, contact form',
     'https://www.broadridge.com/'],
    ['Industry / Who We Serve', 'High',
     'Hero with H1, feature card carousel, accordion solutions, statistics counter, awards section, insights carousel, sub-segment carousel, contact form',
     'https://www.broadridge.com/who-we-serve/asset-management'],
    ['Capability Category', 'Medium',
     'Hero with gradient background, feature carousel, solutions card grid, statistics counter, awards grid, contact form',
     'https://www.broadridge.com/capability/front-office-solutions/'],
    ['Product Detail', 'Medium',
     'Hero with breadcrumb + background image, feature columns with bullet lists, accordion, testimonial quote, related solutions carousel, insights carousel, contact form',
     'https://www.broadridge.com/capability/front-office-solutions/connectivity/'],
    ['Insights / Article', 'Medium',
     'Hero with image, long-form article content, inline quote blocks, image gallery, FAQ accordion, timeline component, related insights, contact form',
     'https://www.broadridge.com/insights/cci-implications-for-individual-firms-the-industry-and-investors'],
    ['Press Release', 'Low',
     'Breadcrumb, title, date, long-form text, sidebar with share buttons and archive link, media contact info. No contact form.',
     'https://www.broadridge.com/press-release/2026/broadridge-to-acquire-cqg'],
    ['About / Corporate', 'Medium',
     'Hero with video background, statistics counter, link card grid (6), CTA banners (careers, IR), partner logo carousel, contact form',
     'https://www.broadridge.com/about/'],
    ['Newsroom', 'Medium',
     'Hero with image, press release list (date + title), featured news cards, news card grid, quick links carousel, media contacts with photos',
     'https://www.broadridge.com/news-room'],
    ['Case Study (Legacy)', 'High',
     'LEGACY CMS: Different navigation, hero block, content with sidebar (share + related resources), different form structure with jQuery, completely different footer. Requires full manual migration.',
     'https://www.broadridge.com/case-study/asset-management/long-short-equity-hedge-fund-enhances-portfolio-management'],
    ['Campaign / Landing Page', 'Low',
     'Hero with large icon/illustration, article text, CTA button to PDF download, contact form',
     'https://www.broadridge.com/campaign/wrapper-to-revolution-etf-share-classes-as-a-structural-shift'],
    ['Leadership Team', 'Medium',
     'Hero with breadcrumb, sticky sidebar navigation, people grid (photo + name + title + bio link), sections for Executive Leadership, Board of Directors, In Memoriam',
     'https://www.broadridge.com/our-leadership-team'],
    ['Contact Us', 'Medium',
     'Hero, general inquiries (3 regional HQs), quick links grid, sticky sidebar region navigation, office listing by country with addresses/phones, contact form',
     'https://www.broadridge.com/contact-us'],
]

add_table_with_data(
    doc,
    ['Template Name', 'Complexity', 'Key Components', 'Reference URL'],
    templates_data,
    col_widths=[1.3, 0.7, 3.2, 1.8]
)

# Add screenshots for each template
doc.add_paragraph('')
doc.add_heading('Template Screenshots', level=2)

screenshot_map = [
    ('screenshots/broadridge-homepage-full.png', 'Figure 2.1: Homepage Template'),
    ('screenshots/broadridge-industry-asset-mgmt.png', 'Figure 2.2: Industry / Who We Serve Template'),
    ('screenshots/broadridge-capability-front-office.png', 'Figure 2.3: Capability Category Template'),
    ('screenshots/broadridge-product-detail-connectivity.png', 'Figure 2.4: Product Detail Template'),
    ('screenshots/broadridge-insights-article.png', 'Figure 2.5: Insights / Article Template'),
    ('screenshots/broadridge-press-release.png', 'Figure 2.6: Press Release Template'),
    ('screenshots/broadridge-about.png', 'Figure 2.7: About / Corporate Template'),
    ('screenshots/broadridge-newsroom.png', 'Figure 2.8: Newsroom Template'),
    ('screenshots/broadridge-case-study.png', 'Figure 2.9: Case Study (Legacy CMS) Template'),
    ('screenshots/broadridge-campaign-landing.png', 'Figure 2.10: Campaign / Landing Page Template'),
    ('screenshots/broadridge-leadership.png', 'Figure 2.11: Leadership Team Template'),
    ('screenshots/broadridge-contact-us.png', 'Figure 2.12: Contact Us Template'),
]

for img_path, caption in screenshot_map:
    full_path = f'/workspace/{img_path}'
    if os.path.exists(full_path):
        add_screenshot(doc, full_path, caption, width=5.0)
        doc.add_paragraph('')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. BLOCKS / COMPONENTS CATALOG
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. Blocks / Components Catalog', level=1)

doc.add_paragraph('The following reusable blocks and components were identified across the site. Design variations of blocks sharing the same content model are listed as variants rather than separate blocks.')

blocks_data = [
    ['Hero Banner', 'High', 'Interactive',
     'Full-width hero with animated background (Lottie), title, subtitle, CTA button. Variants: (a) Lottie animation (homepage), (b) Static background image, (c) Video background (About), (d) Gradient background (Capability)',
     'All pages'],
    ['Navigation / Header', 'High', 'Interactive',
     'Sticky header with logo, mega-menu with multi-level dropdowns, search toggle, contact button, hamburger menu (mobile). Includes locale selector.',
     'All pages'],
    ['Footer', 'Medium', 'Static',
     'Multi-column footer with company info, NYSE stock ticker, social media links, navigation links, locale links, legal links, copyright. Variants: (a) Modern footer, (b) Legacy footer (case study pages)',
     'All pages'],
    ['Contact Form (Flyout)', 'High', 'Interactive',
     'Full-width contact form section with fields: first name, last name, email, phone, job title, company, country (dropdown), message. reCAPTCHA protected. Regional phone numbers sidebar. Variants: (a) Inline flyout form, (b) Legacy form (different fields/layout on case study)',
     'Most pages'],
    ['Card Carousel / Slider', 'High', 'Interactive',
     'Swiper.js-powered carousel with navigation arrows, pagination dots, and swipe support. Variants: (a) Insights card carousel (image + category tag + title), (b) Featured solutions carousel (icon + title + description), (c) Sub-segment carousel, (d) Quick links carousel, (e) Partner logo carousel',
     'Homepage, Industry, Product Detail, About, Newsroom'],
    ['Tabbed Content', 'High', 'Interactive',
     'Alpine.js-powered tabbed navigation with animated content panels. Each tab shows an image, description text, and CTA link. Used for industry segments and capabilities.',
     'Homepage'],
    ['Accordion / Expandable', 'Medium', 'Interactive',
     'Expandable content sections with click-to-toggle panels. Variants: (a) Solutions accordion (Industry pages), (b) FAQ accordion (Article pages), (c) Capabilities accordion (Homepage)',
     'Homepage, Industry, Articles'],
    ['Statistics Counter', 'Medium', 'Static',
     'Horizontal row of large numeric statistics with labels. Typically 3-4 stats displayed. Used on Industry, Capability, and About pages.',
     'Industry, Capability, About'],
    ['Card Grid', 'Medium', 'Static',
     'Responsive grid of linked cards. Variants: (a) Solution cards with icon + title + description (Capability), (b) Link cards with image + title (About), (c) News cards (Newsroom), (d) Award cards with logo + description',
     'Multiple pages'],
    ['Breadcrumb', 'Low', 'Static',
     'Hierarchical breadcrumb trail showing page path. Appears below hero on detail pages.',
     'Product Detail, Press Release, Leadership'],
    ['Testimonial / Quote', 'Low', 'Static',
     'Styled blockquote with attribution. Variants: (a) Full-width testimonial with large quote marks (Product Detail), (b) Inline pull quote within article content',
     'Product Detail, Articles'],
    ['Timeline', 'Medium', 'Static',
     'Vertical timeline with dated entries. Used in article/insight content for chronological information.',
     'Articles/Insights'],
    ['CTA Banner', 'Low', 'Static',
     'Full-width banner with background color/image, headline text, and CTA button. Used for cross-promotion (careers, investor relations, etc.).',
     'About, various pages'],
    ['People Grid', 'Medium', 'Static',
     'Grid of person cards with photo, name, title, and link to bio. Grouped by organizational section with sticky sidebar navigation.',
     'Leadership Team'],
    ['Office Directory', 'Medium', 'Static',
     'Structured listing of office locations grouped by country/region with addresses, phone numbers, and sticky sidebar for navigation.',
     'Contact Us'],
    ['Press Release List', 'Low', 'Static',
     'Chronological list of press releases with date and title link. Simple text-based list format.',
     'Newsroom'],
    ['Share Buttons / Social', 'Low', 'Interactive',
     'Social sharing buttons (LinkedIn, Twitter, email, print) displayed in sidebar on content pages.',
     'Press Release, Case Study, Articles'],
    ['Image Gallery / Carousel', 'Medium', 'Interactive',
     'Inline image carousel within article content. Uses Swiper.js with lightbox capability.',
     'Articles/Insights'],
    ['Cookie Consent Banner', 'Medium', 'Interactive',
     'OneTrust-powered cookie consent dialog with Accept All, Reject All, and Customize Settings options.',
     'All pages (overlay)'],
    ['Search Overlay', 'Medium', 'Interactive',
     'Full-screen search overlay triggered from header. Includes search input with suggestions/autocomplete.',
     'All pages (header)'],
    ['Sidebar Navigation', 'Medium', 'Interactive',
     'Sticky sidebar with section links for in-page navigation. Highlights current section on scroll.',
     'Leadership, Contact Us'],
    ['Video Player', 'Medium', 'Interactive',
     'Embedded video player (appears in hero sections on About page). Auto-play with controls.',
     'About page'],
    ['Lottie Animation', 'High', 'Interactive',
     'Custom Lottie animation rendered via lottie-web library. Used for homepage hero animated illustration.',
     'Homepage'],
    ['Announcement Bar', 'Low', 'Static',
     'Horizontal announcement strip below hero linking to featured content with category label and arrow icon.',
     'Homepage'],
    ['Media Contact Card', 'Low', 'Static',
     'Card with photo, name, title, and contact info for media relations personnel.',
     'Newsroom'],
]

add_table_with_data(
    doc,
    ['Block Name', 'Complexity', 'Behavior', 'Description & Variants', 'Locations'],
    blocks_data,
    col_widths=[1.1, 0.6, 0.6, 3.2, 1.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. PAGE COUNTS BY TEMPLATE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. Page Counts by Template', level=1)

doc.add_paragraph('Page counts are derived from the sitemap (https://www.broadridge.com/us-en/sitemap.xml) which indexes 1,084 URLs. URLs are mapped to template types based on path patterns. Multi-locale pages (fr/, de/, jp/) are not indexed in the primary sitemap and are estimated separately.')

page_counts_data = [
    ['Press Release', '319', '~29%', 'Automated', 'Repeatable template with consistent structure. Ideal for bulk automated migration.'],
    ['Insights / Article', '171', '~16%', 'Automated', 'Consistent article template. May need manual QA for inline components (timelines, galleries).'],
    ['Capability Category + Product Detail', '116', '~11%', 'Semi-Automated', 'Category landing pages are repeatable; product detail pages may have variant layouts.'],
    ['Insight Pages (Hub/Index)', '57', '~5%', 'Semi-Automated', 'Hub pages with dynamic content aggregation. Need manual review for filtering logic.'],
    ['Campaign / Landing', '60', '~6%', 'Semi-Automated', 'Mix of /campaign/ (46) and /campaigns/ (14) paths. Varying layouts require review.'],
    ['Legal Pages', '27', '~2%', 'Automated', 'Simple text-heavy pages with consistent format.'],
    ['About / Corporate', '6', '~1%', 'Manual', 'Each page has unique layout and content combinations.'],
    ['Hub Pages', '6', '~1%', 'Manual', 'Custom hub/landing pages with unique structures.'],
    ['Newsroom', '1', '<1%', 'Manual', 'Single page with dynamic content aggregation (press releases, news cards).'],
    ['Contact Us', '1', '<1%', 'Manual', 'Single page with complex office directory and regional data.'],
    ['Leadership Team', '1', '<1%', 'Manual', 'Single page with people grid and sticky navigation.'],
    ['Homepage', '1', '<1%', 'Manual', 'Unique page with Lottie animations, multiple sections, and high complexity.'],
    ['Case Study (Legacy CMS)', '~15-30 est.', '~2%', 'Manual', 'Legacy CMS platform. Requires full manual migration due to different technology stack.'],
    ['Industry / Who We Serve', '~5', '<1%', 'Semi-Automated', 'Repeatable template but limited page count.'],
    ['Other / Utility', '~4', '<1%', 'Manual', 'CMS reports, page lists, etc. May be excluded from migration.'],
    ['SUBTOTAL (English)', '~1,084', '100%', '—', '—'],
    ['Multi-locale (fr/, de/, jp/)', '~200-400 est.', '—', 'Automated', 'Translated versions of key pages. Can be automated once English templates are established.'],
    ['TOTAL ESTIMATED', '~1,300-1,500', '—', '—', '—'],
]

add_table_with_data(
    doc,
    ['Template Type', 'Page Count', '% of Total', 'Migration Approach', 'Notes'],
    page_counts_data,
    col_widths=[1.3, 0.7, 0.6, 0.9, 2.8]
)

doc.add_paragraph('')
doc.add_heading('Migration Approach Summary', level=2)

approach_data = [
    ['Automated (bulk import)', '~570', '~53%', 'Press releases, articles, legal pages, multi-locale pages'],
    ['Semi-Automated (template + manual QA)', '~240', '~22%', 'Capabilities, campaigns, insight hubs, industry pages'],
    ['Manual', '~75-100', '~8%', 'Homepage, about, newsroom, contact, leadership, legacy case studies, hub pages'],
    ['Excluded / Deferred', '~15-20', '~2%', 'Utility pages, CMS-specific pages, external redirects'],
    ['Multi-locale (post-English)', '~200-400', '~15%', 'Translated versions after English templates established'],
]

add_table_with_data(
    doc,
    ['Approach', 'Est. Pages', '% of Total', 'Page Types'],
    approach_data,
    col_widths=[1.5, 0.8, 0.7, 3.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. INTEGRATIONS ANALYSIS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. Integrations Analysis', level=1)

doc.add_paragraph('The following third-party integrations were identified through analysis of script tags, external domains, iframes, and meta tags loaded on Broadridge.com pages.')

integrations_data = [
    ['Google Tag Manager', 'Tag Management', 'High',
     'Container ID: GTM-PW7DJ8. Multiple GA4 properties: G-C8WF5M11DE, G-C83WJMZCP3, G-3PFPJ427KS. Google Ads conversion tracking (AW-869648502).',
     'All pages'],
    ['OneTrust / CookieLaw', 'Cookie Consent (Privacy)', 'High',
     'Cookie consent management platform (cdn.cookielaw.org). SDK stub + banner. Categories: C0001-C0004. "Do Not Sell" link in footer.',
     'All pages'],
    ['Google reCAPTCHA v3', 'Form Security', 'Medium',
     'Invisible reCAPTCHA (site key: 6Ld5roIrAAAAAC4ZCTvZTEBVv30JkZouN4To_bCW). Protects contact forms from spam submissions.',
     'Pages with contact form'],
    ['Facebook Pixel', 'Advertising / Analytics', 'Medium',
     'Pixel ID: 1184250882355912. Tracks page views and conversions for Facebook ad campaigns.',
     'All pages'],
    ['LinkedIn Insight Tag', 'Advertising / Analytics', 'Medium',
     'LinkedIn analytics for conversion tracking and audience building (snap.licdn.com).',
     'All pages'],
    ['Microsoft Clarity', 'Session Recording / Heatmaps', 'Medium',
     'User behavior analytics with session recordings and heatmaps (scripts.clarity.ms).',
     'All pages'],
    ['Bing UET Tag', 'Advertising / Analytics', 'Medium',
     'Microsoft Advertising Universal Event Tracking. Tag ID: 25073696 (bat.bing.com).',
     'All pages'],
    ['ZoomInfo', 'Visitor Intelligence', 'Medium',
     'B2B visitor identification and intent data (js.zi-scripts.com/zi-tag.js). Identifies company-level website visitors.',
     'All pages'],
    ['The Trade Desk', 'Advertising', 'Low',
     'Programmatic advertising pixel. Advertiser ID: 7ts8ctf. Universal pixel for conversion tracking (js.adsrvr.org).',
     'All pages'],
    ['Clickagy', 'Audience Intelligence', 'Low',
     'Intent data and audience intelligence platform (tags.clickagy.com). B2B targeting.',
     'All pages'],
    ['Pardot / Marketing Cloud', 'Marketing Automation', 'High',
     'Salesforce Pardot tracking (pages.broadridge.com). Account ID: 1071212. Form handling and lead tracking. Also powers email preference management.',
     'All pages + forms'],
    ['CrownPeak (Omni)', 'Accessibility Monitoring', 'Medium',
     'Digital accessibility monitoring and compliance (snippet.omm.crownpeak.com). Ensures WCAG compliance.',
     'All pages'],
    ['Google DoubleClick', 'Advertising', 'Low',
     'Google advertising network for remarketing and display ads (googleads.g.doubleclick.net).',
     'All pages'],
    ['Lottie (lottie-web)', 'Animation Library', 'Medium',
     'CDN-hosted animation library (cdnjs.cloudflare.com/lottie-web/5.12.2). Powers homepage hero animation.',
     'Homepage'],
    ['ipify API', 'IP Geolocation', 'Low',
     'IP address detection service (api.ipify.org) used by geolocation redirect script for locale-based routing.',
     'All pages'],
    ['Workday', 'HR / Careers (External)', 'Low',
     'External link to Workday career portal (broadridge.wd5.myworkdayjobs.com). Not embedded, but linked from footer and About pages.',
     'Footer, About'],
]

add_table_with_data(
    doc,
    ['Integration', 'Type', 'Complexity', 'Details', 'Scope'],
    integrations_data,
    col_widths=[1.1, 0.9, 0.6, 3.0, 0.9]
)

doc.add_paragraph('')
doc.add_heading('Integration Migration Considerations', level=2)

int_considerations = """
• Google Tag Manager: Can be re-implemented via a single GTM container script in the EDS head.html. All downstream tags (GA4, Google Ads, Facebook Pixel, LinkedIn, Bing, etc.) are managed through GTM and should transfer with minimal effort.

• OneTrust: Requires the OneTrust SDK script in head.html with the same domain configuration. Cookie categories (C0001-C0004) must be preserved to maintain consent preferences. The "Do Not Sell" link must be maintained in the footer.

• reCAPTCHA: Form submissions will need reCAPTCHA integration. This requires the API script and server-side validation. Consider using an EDS form service or custom form block.

• Pardot/Marketing Cloud: The most complex integration. Form submissions currently post to pages.broadridge.com. The Pardot tracking script and form handlers need careful migration. Consider implementing via a custom form block that submits to the Pardot endpoint.

• ZoomInfo & Visitor Intelligence: Simple script tags that can be loaded via GTM or head.html.

• CrownPeak Accessibility: Will need to be re-validated after migration to ensure WCAG compliance with EDS markup.

• Geolocation Redirect: The custom geolocation.js script uses ipify for IP detection and redirects users to locale-specific paths. This logic needs to be replicated in EDS, potentially via a custom script or Edge worker.
"""
doc.add_paragraph(int_considerations.strip())

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. COMPLEX USE CASES & OBSERVATIONS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. Complex Use Cases & Observations', level=1)

doc.add_paragraph('The following edge cases, complex workflows, and special considerations were identified during the site analysis.')

complex_data = [
    ['Legacy CMS Pages (Case Studies)', 'High', '~15-30 pages',
     '/case-study/* URLs',
     'These pages run on a completely different CMS platform with: (1) Different header/navigation, (2) Different footer, (3) Different form structure (more fields, different endpoint), (4) jQuery-based interactions instead of Alpine.js, (5) Different CSS framework. These CANNOT be migrated using the same templates as modern pages and require individual manual migration.'],
    ['Lottie Hero Animation', 'High', '1 page',
     'Homepage',
     'The homepage hero uses a custom Lottie animation (lottie-web library + custom JS) with animated SVG illustration. This requires a custom EDS block with the Lottie library loaded and the animation JSON file served correctly.'],
    ['Multi-Locale Geolocation Redirect', 'High', '~4 locales',
     'All pages (geolocation.js)',
     'A JavaScript-based geolocation redirect system detects user IP via ipify API and redirects to locale-specific paths (/fr/, /de/, /jp/). This needs careful re-implementation in EDS, potentially using Edge functions or a redirect service.'],
    ['Swiper.js Carousels', 'Medium', '~20+ instances',
     'Multiple pages',
     'Extensive use of Swiper.js for card carousels, image galleries, and content sliders. Each carousel variant has different configurations (slides per view, pagination type, autoplay). Requires a flexible carousel block in EDS.'],
    ['Alpine.js Interactive Components', 'Medium', '~10+ instances',
     'Homepage, various pages',
     'Alpine.js powers tabbed content, accordion sections, and form interactions. These declarative behaviors need to be re-implemented as EDS block JavaScript decorations.'],
    ['NYSE Stock Ticker (Footer)', 'Medium', '1 instance',
     'Footer (all pages)',
     'The footer displays live Broadridge (BR) stock price from NYSE with change percentage. This requires a real-time or cached data feed integration in the EDS footer.'],
    ['Contact Form Regional Logic', 'Medium', '~50+ pages',
     'Most pages with flyout form',
     'The contact form includes a country dropdown that may influence form routing. Regional phone numbers are displayed alongside the form. reCAPTCHA validation and Pardot submission endpoint must be maintained.'],
    ['Pardot Form Tracking', 'Medium', 'All pages',
     'pages.broadridge.com integration',
     'Pardot (Salesforce) tracking is loaded on every page for visitor analytics and form handling. The account ID (1071212) and visitor tracking must be preserved. Form submissions use Pardot endpoints.'],
    ['Dynamic Content Aggregation', 'Medium', '~57+ pages',
     '/insight-pages/* URLs',
     'Insight hub pages aggregate and filter content dynamically (articles, whitepapers, case studies). This likely requires a custom EDS block with client-side data fetching or a spreadsheet-based index.'],
    ['Video Backgrounds', 'Medium', '1-2 pages',
     'About page',
     'The About page hero includes a looping video background. Requires a video-capable hero block with proper autoplay, mute, and fallback image handling.'],
    ['Sticky Sidebar Navigation', 'Medium', '2-3 pages',
     'Leadership, Contact Us',
     'Some pages use sticky sidebar navigation that highlights the active section on scroll. Requires a custom EDS block with IntersectionObserver for scroll tracking.'],
    ['PDF Downloads / Gated Content', 'Low', '~20+ instances',
     'Campaign pages, insights',
     'Several campaign and insight pages link to PDF downloads, some potentially gated behind form submissions. Migration needs to handle PDF asset management and gating logic.'],
    ['External Subdomain Links', 'Low', 'Multiple',
     'Footer, About, various',
     'Links to external Broadridge properties: broadridge-ir.com (Investor Relations), pages.broadridge.com (Email Preferences), broadridge.wd5.myworkdayjobs.com (Careers). These are external links, not pages to migrate, but must be preserved.'],
    ['Open Graph / Social Metadata', 'Low', 'All pages',
     'Meta tags',
     'All pages include comprehensive OG and Twitter Card metadata (title, description, image). This must be generated correctly in EDS page metadata.'],
    ['WCAG Accessibility Compliance', 'Medium', 'All pages',
     'CrownPeak + custom a11y.js',
     'The site uses CrownPeak accessibility monitoring and has a custom accessibility.js script. EDS migration must maintain WCAG 2.1 AA compliance. Skip-to-content link, ARIA labels, and keyboard navigation must be preserved.'],
]

add_table_with_data(
    doc,
    ['Use Case', 'Complexity', 'Instance Count', 'Location(s)', 'Details & Considerations'],
    complex_data,
    col_widths=[1.2, 0.6, 0.7, 1.0, 3.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. MIGRATION ESTIMATES
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. Migration Estimates', level=1)

doc.add_paragraph('The following estimates are based on the analysis of site templates, blocks, integrations, and page complexity. Estimates assume a team with Edge Delivery Services experience.')

doc.add_heading('7.1 Phase Breakdown', level=2)

phases_data = [
    ['Phase 1: Discovery & Design System', '',
     'Extract design tokens (colors, typography, spacing), establish EDS project structure, configure styles.css, set up development environment.',
     ''],
    ['  - Design token extraction', '1-2 weeks', '', 'Global colors, fonts, spacing, breakpoints'],
    ['  - EDS project setup', '1 week', '', 'Boilerplate, styles.css, header/footer shells'],
    ['  - Navigation structure', '1 week', '', 'Mega-menu, mobile nav, locale selector'],
    ['Phase 2: Block Development', '',
     'Develop all 25 reusable blocks/components identified in the catalog.',
     ''],
    ['  - High complexity blocks (6)', '3-4 weeks', '', 'Hero (4 variants), nav, contact form, carousel, tabs, Lottie'],
    ['  - Medium complexity blocks (11)', '3-4 weeks', '', 'Accordion, stats counter, card grid, timeline, sidebar nav, etc.'],
    ['  - Low complexity blocks (8)', '1-2 weeks', '', 'Breadcrumb, CTA banner, quote, share buttons, etc.'],
    ['Phase 3: Template Migration', '',
     'Create template pages for each of the 12 identified templates.',
     ''],
    ['  - Template creation (12 templates)', '2-3 weeks', '', 'One representative page per template with all blocks'],
    ['  - Template QA & refinement', '1-2 weeks', '', 'Visual comparison, responsive testing, accessibility'],
    ['Phase 4: Content Migration', '',
     'Migrate all ~1,084 English pages using automated and manual approaches.',
     ''],
    ['  - Automated migration (press releases, articles, legal)', '2-3 weeks', '', '~570 pages via import scripts'],
    ['  - Semi-automated migration (capabilities, campaigns)', '2-3 weeks', '', '~240 pages with template + manual QA'],
    ['  - Manual migration (complex/unique pages)', '2-3 weeks', '', '~75-100 pages including legacy CMS'],
    ['  - Legacy CMS pages (case studies)', '1-2 weeks', '', '~15-30 pages requiring full manual rebuild'],
    ['Phase 5: Multi-Locale', '',
     'Migrate French, German, and Japanese locale pages.',
     ''],
    ['  - Locale configuration', '1 week', '', 'Multi-site setup, locale routing'],
    ['  - Content migration (3 locales)', '2-4 weeks', '', '~200-400 pages across 3 languages'],
    ['Phase 6: Integrations', '',
     'Re-implement all third-party integrations.',
     ''],
    ['  - GTM + analytics stack', '1 week', '', 'GTM container, GA4, Facebook, LinkedIn, Bing, etc.'],
    ['  - OneTrust cookie consent', '1 week', '', 'Cookie banner, preference center, consent categories'],
    ['  - Forms + Pardot + reCAPTCHA', '2 weeks', '', 'Contact form block with Pardot submission and reCAPTCHA'],
    ['  - Geolocation redirect', '1 week', '', 'IP-based locale redirect logic'],
    ['  - Other (ZoomInfo, CrownPeak, etc.)', '1 week', '', 'Simple script inclusions'],
    ['Phase 7: QA & Launch', '',
     'Comprehensive testing, UAT, and launch preparation.',
     ''],
    ['  - Automated visual regression', '1-2 weeks', '', 'Page-by-page visual comparison'],
    ['  - Manual QA & UAT', '2-3 weeks', '', 'Cross-browser, responsive, accessibility, forms'],
    ['  - Performance optimization', '1 week', '', 'Core Web Vitals, Lighthouse scoring'],
    ['  - Launch preparation', '1 week', '', 'DNS, redirects, monitoring, rollback plan'],
]

add_table_with_data(
    doc,
    ['Phase / Task', 'Duration', 'Description', 'Notes'],
    phases_data,
    col_widths=[2.0, 0.8, 2.5, 1.5]
)

doc.add_paragraph('')
doc.add_heading('7.2 Overall Timeline Summary', level=2)

timeline_data = [
    ['Phase 1: Discovery & Design System', '3-4 weeks'],
    ['Phase 2: Block Development', '7-10 weeks'],
    ['Phase 3: Template Migration', '3-5 weeks'],
    ['Phase 4: Content Migration (English)', '7-11 weeks'],
    ['Phase 5: Multi-Locale Migration', '3-5 weeks'],
    ['Phase 6: Integrations', '5-6 weeks'],
    ['Phase 7: QA & Launch', '5-7 weeks'],
    ['TOTAL (Sequential)', '33-48 weeks'],
    ['TOTAL (With Parallelization)', '20-30 weeks'],
]

add_table_with_data(
    doc,
    ['Phase', 'Estimated Duration'],
    timeline_data,
    col_widths=[3.5, 2.0]
)

doc.add_paragraph('')
doc.add_paragraph('Note: Phases 2-3 can partially overlap. Phases 4-6 can run in parallel. The parallelized estimate assumes a team of 4-6 developers working concurrently across streams.')

doc.add_heading('7.3 Team & Effort Estimate', level=2)

effort_data = [
    ['EDS Architect / Tech Lead', '1', 'Full project', 'Architecture decisions, block design, code reviews, integration planning'],
    ['EDS Developer (Blocks)', '2-3', 'Phases 2-3', 'Block JavaScript/CSS development, template creation'],
    ['Content Migration Specialist', '1-2', 'Phases 4-5', 'Import scripts, content mapping, QA'],
    ['Design / CSS Specialist', '1', 'Phases 1-3', 'Design tokens, global styles, responsive CSS'],
    ['QA Engineer', '1-2', 'Phase 7 + ongoing', 'Visual regression, accessibility, cross-browser testing'],
    ['Project Manager', '1', 'Full project', 'Coordination, stakeholder communication, timeline management'],
]

add_table_with_data(
    doc,
    ['Role', 'Headcount', 'Duration', 'Responsibilities'],
    effort_data,
    col_widths=[1.5, 0.7, 1.0, 3.5]
)

doc.add_paragraph('')
doc.add_heading('7.4 Risk Factors', level=2)

risks = """
1. Legacy CMS Pages: The case study pages running on a different CMS represent the highest risk item. If more legacy pages are discovered beyond the case study section, the manual migration effort will increase significantly.

2. Form / Pardot Integration: The contact form with reCAPTCHA and Pardot submission is present on most pages. Any issues with form migration will have wide-reaching impact. Early prototyping of the form block is recommended.

3. Multi-Locale Scope: The exact number of localized pages is uncertain (200-400 estimated). A full audit of /fr/, /de/, and /jp/ content is recommended before Phase 5 begins.

4. Swiper.js Carousel Complexity: The site uses multiple Swiper.js carousel variants with different configurations. Each variant requires careful development and testing. Consider using an EDS carousel plugin or library block.

5. Dynamic Content Aggregation: Insight hub pages (/insight-pages/) appear to aggregate content dynamically. Implementing equivalent functionality in EDS may require spreadsheet-based indexes and custom filtering blocks.

6. SEO Continuity: With 1,084+ indexed pages, maintaining URL structures, metadata, canonical tags, and redirect mappings is critical to preserve organic search rankings.

7. Accessibility Compliance: The site currently has CrownPeak accessibility monitoring. EDS output markup will differ from the current CMS, requiring fresh WCAG 2.1 AA validation.
"""
doc.add_paragraph(risks.strip())

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8. APPENDIX
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. Appendix', level=1)

doc.add_heading('8.1 Technology Stack Summary', level=2)

tech_data = [
    ['CMS Platform', 'Custom CMS ("Broadridge V2") - Modern pages\nLegacy CMS - Case study pages'],
    ['JavaScript Frameworks', 'jQuery 3.7.1, Alpine.js (bundled), Swiper.js (bundled), Lottie-web 5.12.2'],
    ['CSS', 'Custom CSS via styles.css, Adobe Typekit fonts (xmw3hcn)'],
    ['Tag Management', 'Google Tag Manager (GTM-PW7DJ8)'],
    ['Analytics', 'Google Analytics 4 (3 properties), Facebook Pixel, LinkedIn Insight, Microsoft Clarity, Bing UET'],
    ['Privacy', 'OneTrust CookieLaw (cookie consent banner)'],
    ['Forms', 'Custom forms with reCAPTCHA v3 + Pardot/Marketing Cloud backend'],
    ['Marketing Automation', 'Salesforce Pardot (pages.broadridge.com, Account: 1071212)'],
    ['Visitor Intelligence', 'ZoomInfo (zi-tag.js), Clickagy'],
    ['Advertising', 'Google Ads, The Trade Desk, DoubleClick'],
    ['Accessibility', 'CrownPeak Omni + custom accessibility.js'],
    ['CDN / Hosting', 'Custom hosting (broadridge.com), Cloudflare CDN for libraries'],
    ['Locales', 'English (default), French (fr/), German (de/), Japanese (jp/)'],
    ['IP Detection', 'ipify API for geolocation-based locale routing'],
]

add_table_with_data(
    doc,
    ['Component', 'Details'],
    tech_data,
    col_widths=[1.5, 5.0]
)

doc.add_paragraph('')
doc.add_heading('8.2 Sitemap Statistics', level=2)

sitemap_data = [
    ['Total indexed URLs (English)', '1,084'],
    ['Estimated multi-locale pages', '200-400'],
    ['Press releases', '319 (29%)'],
    ['Insights / Articles', '171 (16%)'],
    ['Capabilities', '116 (11%)'],
    ['Insight Hub Pages', '57 (5%)'],
    ['Campaigns', '60 (6%)'],
    ['Legal', '27 (2%)'],
    ['Other (About, Newsroom, etc.)', '~28 (3%)'],
    ['Estimated legacy CMS pages', '15-30'],
]

add_table_with_data(
    doc,
    ['Metric', 'Count'],
    sitemap_data,
    col_widths=[3.0, 2.0]
)

doc.add_paragraph('')
doc.add_heading('8.3 External Domains Loaded', level=2)

domains_data = [
    ['js.adsrvr.org', 'The Trade Desk (advertising)'],
    ['tags.clickagy.com', 'Clickagy (audience intelligence)'],
    ['scripts.clarity.ms', 'Microsoft Clarity (session recording)'],
    ['www.googletagmanager.com', 'Google Tag Manager + GA4'],
    ['www.gstatic.com', 'Google reCAPTCHA'],
    ['snap.licdn.com', 'LinkedIn Insight Tag'],
    ['connect.facebook.net', 'Facebook Pixel'],
    ['bat.bing.com', 'Bing UET Tag'],
    ['cdn.cookielaw.org', 'OneTrust CookieLaw'],
    ['snippet.omm.crownpeak.com', 'CrownPeak Accessibility'],
    ['www.google.com', 'Google reCAPTCHA API'],
    ['googleads.g.doubleclick.net', 'Google DoubleClick (advertising)'],
    ['cdnjs.cloudflare.com', 'Lottie-web library (CDN)'],
    ['api.ipify.org', 'IP detection for geolocation'],
    ['js.zi-scripts.com', 'ZoomInfo visitor intelligence'],
    ['pages.broadridge.com', 'Pardot / Salesforce Marketing Cloud'],
]

add_table_with_data(
    doc,
    ['Domain', 'Purpose'],
    domains_data,
    col_widths=[2.5, 4.0]
)

# ── Save ────────────────────────────────────────────────────────
output_path = '/workspace/Broadridge_Site_Analysis_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB')
